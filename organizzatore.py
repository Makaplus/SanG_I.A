#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
codex/iniziare-progetto-libreria-atti-di-polizia-sri8es
SANGIA - ORGANIZZATORE (v21.2)

Fix/novità principali:
1) Pre-process per estensione:
   - .doc/.docx/.rtf/.txt: chiama SEMPRE organizzatore_doc_reader.py (anche se non è OCC)
   - .pdf: preview pypdf per classificazione (OCR forte resta ai worker/Router revisione)

2) Classificazione rigida via “tabella regole” nel DB:
   - classi (attuali): OCC -> INFORMATIVA -> RELAZIONE -> ALTRO
   - seed DB include anche placeholder: SENTENZA, ARTICOLI, ANNOTAZIONI, RICORSI (per step successivi)

3) Smistamento NON-OCC:
   - se anno noto nel DB -> usa organizzatore_percorso.py (sposta in libreria\YYYY\) senza rename
   - se anno mancante -> usa organizzatore_vari_verbali.py

4) Output: stampa SEMPRE "tipo=..." su ogni riga.
"""

import re
import time
import json
import hashlib
import sqlite3
import subprocess
import sys
from pathlib import Path
from typing import Dict, Any, List, Optional, Tuple

# =============================================================================
# UTILS
# =============================================================================

def keep_last_n_files(folder: Path, n: int = 5):
    folder.mkdir(parents=True, exist_ok=True)
    files = [p for p in folder.iterdir() if p.is_file()]
    files.sort(key=lambda p: p.stat().st_mtime, reverse=True)
    for p in files[n:]:
        try:
            p.unlink()
        except Exception:
            pass

def sha1_file(path: Path) -> str:
    h = hashlib.sha1()
    with open(path, "rb") as f:
        while True:
            chunk = f.read(1024 * 1024)
            if not chunk:
                break
            h.update(chunk)
    return h.hexdigest()

def format_mtime(ts: float) -> str:
    return time.strftime("%Y-%m-%d %H:%M:%S", time.localtime(ts))

def mark(ok: bool) -> str:
    return "✅" if ok else "⚠️"

def fmt_tipo(tipo_doc: Optional[str]) -> str:
    t = (tipo_doc or "").strip().upper() or "ALTRO"
    return f"tipo={t}"

def filename_has_occ_token(name: str) -> bool:
    # token OCC/OCCC come parola (evita match su 'OCCASIONE' ecc.)
    stem = Path(name).stem
    return bool(re.search(r"(?i)(?:^|[ \._\-])OCC{1,3}(?:$|[ \._\-])", stem))


# =============================================================================
# PATHS
# =============================================================================

PROJECT_ROOT = Path(__file__).resolve().parent
INPUT_DIR = PROJECT_ROOT / "input_documenti"
DUP_DIR = INPUT_DIR / "duplicati"
DB_PATH = PROJECT_ROOT / "libreria" / "documenti.db"

LIB_DIR = PROJECT_ROOT / "libreria"
REV_OCC_DIR = LIB_DIR / "revisione" / "occ"
LOG_DIR = PROJECT_ROOT / "Backup" / "log_smistamenti"

ALLOWED_EXTS = {".pdf", ".docx", ".doc", ".txt", ".rtf"}
MAX_PREVIEW_PAGES_PDF = 2
MAX_PREVIEW_CHARS = 8000
FINAL_STATUSES = {"STORED", "COMPLETED", "DONE"}

# gating OCC completo
RE_RGNR_PARSE = re.compile(r"(\d{1,7})\s*/\s*(\d{4})")

# =============================================================================
# PDF READER
# =============================================================================

def load_pdf_reader():
    try:
        from pypdf import PdfReader  # type: ignore
        return PdfReader, "pypdf"
    except Exception:
        return None, None

PdfReader, PDF_LIB = load_pdf_reader()

def pdf_preview(path: Path) -> Tuple[Dict[str, Any], Optional[int], str]:
    if not PdfReader:
        return {"_pdf_lib": "NONE"}, None, ""
    meta: Dict[str, Any] = {"_pdf_lib": PDF_LIB}
    pages = None
    parts: List[str] = []
    try:
        reader = PdfReader(str(path))
        try:
            pages = len(reader.pages)
        except Exception:
            pages = None

        n = min(MAX_PREVIEW_PAGES_PDF, pages or MAX_PREVIEW_PAGES_PDF)
        for i in range(n):
            try:
                t = reader.pages[i].extract_text() or ""
                t = t.strip()
                if t:
                    parts.append(t)
            except Exception:
                continue
    except Exception as e:
        meta["_error"] = str(e)

    preview = "\n\n".join(parts)
    if len(preview) > MAX_PREVIEW_CHARS:
        preview = preview[:MAX_PREVIEW_CHARS] + "\n...[TRUNCATED]..."
    return meta, pages, preview

def get_preview(path: Path) -> Tuple[Dict[str, Any], Optional[int], str]:
    if path.suffix.lower() == ".pdf":
        return pdf_preview(path)
    return {"_note": "no_preview"}, None, ""

# =============================================================================
# DB / RULES
# =============================================================================

def ensure_db():
    DB_PATH.parent.mkdir(parents=True, exist_ok=True)
    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()

    cur.execute("""
    CREATE TABLE IF NOT EXISTS documenti (
        id INTEGER PRIMARY KEY,
        nome_file TEXT,
        sha1 TEXT,
        dimensione_file INTEGER,
        data_ultima_modifica TEXT,
        numero_pagine INTEGER,
        tipo_documento TEXT,
        rgnr TEXT,
        anno INTEGER,
        procura TEXT,
        status TEXT,
        evidence_occ TEXT,
        conf_occ REAL,
        snippet_testo TEXT,
        data_inserimento TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )
    """)

    cur.execute("PRAGMA table_info(documenti)")
    cols = {row[1] for row in cur.fetchall()}

    def add_col(name: str, coltype: str):
        if name not in cols:
            cur.execute(f"ALTER TABLE documenti ADD COLUMN {name} {coltype}")

    add_col("tipo_file", "TEXT")
    add_col("text_preview", "TEXT")
    add_col("pdf_meta_json", "TEXT")
    add_col("nome_file_orig", "TEXT")
    add_col("nome_file_prev", "TEXT")
    add_col("rename_evidence", "TEXT")
    add_col("percorso_file", "TEXT")
    add_col("percorso_prev", "TEXT")
    add_col("percorso_evidence", "TEXT")
    add_col("dda_flag", "INTEGER")
    add_col("operazione_nome", "TEXT")
    add_col("nome_indagine", "TEXT")
    add_col("forza_polizia", "TEXT")
    add_col("categoria_secondaria", "TEXT")

    add_col("is_scan", "INTEGER")
    add_col("has_rgnr_hint", "INTEGER")
    add_col("rgnr_hint", "TEXT")
    add_col("anno_hint", "INTEGER")
    add_col("hint_text", "TEXT")

    add_col("motivo_revisione", "TEXT")
    add_col("retry_count", "INTEGER")
    add_col("last_retry_at", "TEXT")
    add_col("manual_required", "INTEGER")
    add_col("manual_note", "TEXT")
    add_col("esito_revisione", "TEXT")

    # classificazione rigida (nuovo)
    add_col("class_score", "INTEGER")
    add_col("class_trace", "TEXT")
    add_col("class_forced", "TEXT")
    add_col("note", "TEXT")

    cur.execute("CREATE INDEX IF NOT EXISTS idx_doc_sha1 ON documenti(sha1)")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_doc_status ON documenti(status)")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_doc_tipo ON documenti(tipo_documento)")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_doc_cat2 ON documenti(categoria_secondaria)")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_doc_manual_required ON documenti(manual_required)")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_doc_retry_count ON documenti(retry_count)")

    # tabella regole (minima)
    cur.execute("""
    CREATE TABLE IF NOT EXISTS regole_classificazione (
        id INTEGER PRIMARY KEY,
        classe TEXT NOT NULL,
        marker TEXT NOT NULL,
        peso INTEGER NOT NULL DEFAULT 1,
        tipo_marker TEXT NOT NULL DEFAULT 'forte'  -- forte/debole/blacklist
    )
    """)

    con.commit()
    con.close()


def seed_rules_if_empty():
    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()
    cur.execute("SELECT COUNT(*) FROM regole_classificazione")
    n = cur.fetchone()[0]
    if n and n > 0:
        con.close()
        return

    rules = [
        # OCC (forti)
        ("OCC", "ordinanza di applicazione di misura cautelare", 6, "forte"),
        ("OCC", "ordinanza di custodia cautelare", 6, "forte"),
        ("OCC", "richiesta per l’applicazione di misure cautelari", 7, "forte"),
        ("OCC", "richiesta per l'applicazione di misure cautelari", 7, "forte"),
        ("OCC", "custodia cautelare", 4, "forte"),
        ("OCC", "misura cautelare", 3, "debole"),
        ("OCC", "giudice per le indagini preliminari", 3, "debole"),
        ("OCC", "gip", 2, "debole"),

        # INFORMATIVA (forti)
        ("INFORMATIVA", "comunicazione di notizia di reato", 7, "forte"),
        ("INFORMATIVA", "informativa di reato", 5, "forte"),
        ("INFORMATIVA", "notizia di reato", 4, "debole"),

        # RELAZIONE / ANNOTAZIONI (placeholder per step successivi)
        ("RELAZIONE", "relazione di servizio", 5, "forte"),
        ("ANNOTAZIONE", "annotazione", 4, "forte"),

        # Blacklist OCC (esempi tipici “non atto cautelare”)
        ("OCC", "camera dei deputati", -10, "blacklist"),
        ("OCC", "senato della repubblica", -10, "blacklist"),
        ("OCC", "disegni di legge", -10, "blacklist"),
        ("OCC", "rassegna stampa", -8, "blacklist"),
    ]

    cur.executemany(
        "INSERT INTO regole_classificazione (classe, marker, peso, tipo_marker) VALUES (?,?,?,?)",
        rules
    )
    con.commit()
    con.close()


def fetch_doc_by_sha1(sha1: str):
    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()
    cur.execute("""
        SELECT id, nome_file, tipo_documento, rgnr, anno, procura, dda_flag,
               operazione_nome, nome_indagine, status, percorso_file, forza_polizia,
               text_preview, snippet_testo,
               motivo_revisione, retry_count, manual_required,
               class_score, class_trace, class_forced, note
        FROM documenti
        WHERE sha1=?
        LIMIT 1
    """, (sha1,))
    row = cur.fetchone()
    con.close()
    return row


def db_update_identity(doc_id: int, new_name: str, new_rel_path: str, file_size: int, mtime_str: str):
    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()
    cur.execute("SELECT nome_file, percorso_file FROM documenti WHERE id=? LIMIT 1", (doc_id,))
    row = cur.fetchone()
    old_name = row[0] if row else None
    old_path = row[1] if row else None

    updates = {
        "dimensione_file": file_size,
        "data_ultima_modifica": mtime_str,
    }
    if old_name != new_name:
        updates["nome_file_prev"] = old_name
        updates["nome_file"] = new_name
    if old_path != new_rel_path:
        updates["percorso_prev"] = old_path
        updates["percorso_file"] = new_rel_path

    set_clause = ", ".join([f"{k}=?" for k in updates.keys()])
    params = list(updates.values()) + [doc_id]
    cur.execute(f"UPDATE documenti SET {set_clause} WHERE id=?", params)
    con.commit()
    con.close()


def is_final_status(status: Optional[str]) -> bool:
    return (status or "").upper().strip() in FINAL_STATUSES


def is_path_in_input(percorso_file: Optional[str]) -> bool:
    if not percorso_file:
        return True
    p = percorso_file.replace("/", "\\").lower()
    return p.startswith("input_documenti\\") or p == "input_documenti"


def unique_path(target: Path) -> Path:
    if not target.exists():
        return target
    parent = target.parent
    stem = target.stem
    suf = target.suffix
    for i in range(1, 1000):
        cand = parent / f"{stem} ({i:02d}){suf}"
        if not cand.exists():
            return cand
    raise RuntimeError("Troppi conflitti nome.")


def move_to_duplicati(file_path: Path) -> Path:
    DUP_DIR.mkdir(parents=True, exist_ok=True)
    target = unique_path(DUP_DIR / file_path.name)
    file_path.rename(target)
    return target


def move_occ_to_revisione(file_path: Path, sha1: str) -> Tuple[bool, str]:
    REV_OCC_DIR.mkdir(parents=True, exist_ok=True)
    target = unique_path(REV_OCC_DIR / file_path.name)
    try:
        file_path.rename(target)
    except Exception as e:
        return False, f"move_revisione_fail:{e}"

    new_rel = str(target.relative_to(PROJECT_ROOT)).replace("/", "\\")
    old_rel = str(Path("input_documenti") / file_path.name).replace("/", "\\")

    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()
    cur.execute("SELECT id, nome_file, percorso_file FROM documenti WHERE sha1=? LIMIT 1", (sha1,))
    row = cur.fetchone()
    if not row:
        con.close()
        return False, "db_missing_for_revisione"

    doc_id, old_name, old_path = row

    cur.execute("""
        UPDATE documenti
        SET status='REVISIONE_OCC',
            tipo_documento='OCC',
            percorso_prev=?,
            percorso_file=?,
            nome_file_prev=?,
            nome_file=?,
            motivo_revisione=COALESCE(motivo_revisione, 'MISSING_FIELDS')
        WHERE id=?
    """, (old_path or old_rel, new_rel, old_name, target.name, doc_id))

    con.commit()
    con.close()
    return True, f"moved_to:{new_rel}"


def run_worker(script_rel: Path, args: List[str]) -> Tuple[bool, str]:
    script = (PROJECT_ROOT / script_rel).resolve()
    if not script.exists():
        return False, f"worker_missing:{script}"
    cmd = [sys.executable, str(script)] + args
    proc = subprocess.run(cmd, cwd=str(PROJECT_ROOT), capture_output=True, text=True)
    if proc.returncode != 0:
        return False, (proc.stderr or proc.stdout or "worker_failed").strip()
    return True, (proc.stdout or "").strip()


def iter_input_files() -> List[Path]:
    if not INPUT_DIR.exists():
        return []
    out = []
    for p in INPUT_DIR.iterdir():
        if not p.is_file():
            continue
        if p.parent.name.lower() == "duplicati":
            continue
        if p.suffix.lower() in ALLOWED_EXTS:
            out.append(p)
    return sorted(out)


def db_insert_new_file(f: Path, sha1: str, rel_path: str) -> None:
    meta, pages, preview = get_preview(f)

    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()

    # default neutro, poi classifichiamo con regole
    tipo_doc = "ALTRO"
    status = "NEEDS_DEEPER_ANALYSIS"

    cur.execute("""
        INSERT INTO documenti
        (nome_file, sha1, dimensione_file, data_ultima_modifica, numero_pagine,
         tipo_documento, rgnr, anno, procura, status,
         evidence_occ, conf_occ, snippet_testo,
         tipo_file, text_preview, pdf_meta_json,
         nome_file_orig, percorso_file,
         categoria_secondaria)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (
        f.name, sha1, f.stat().st_size, format_mtime(f.stat().st_mtime), pages,
        tipo_doc, None, None, None, status,
        None, None,
        (preview or "")[:500],
        f.suffix.lower().lstrip("."),
        preview,
        json.dumps(meta, ensure_ascii=False),
        f.name,
        rel_path,
        None
    ))
    con.commit()
    con.close()


def db_update_classification(doc_id: int, tipo: str, status: str, score: int, trace: str, forced: Optional[str]):
    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()
    cur.execute("""
        UPDATE documenti
        SET tipo_documento=?, status=?, class_score=?, class_trace=?, class_forced=?
        WHERE id=?
    """, (tipo, status, int(score), trace, forced, doc_id))
    con.commit()
    con.close()


def score_with_rules(preview_text: str, filename: str) -> Tuple[str, int, str, Optional[str]]:
    """
    Applica regole dal DB e decide una classe.
    Ritorna: (classe, score, trace, forced)
    """
    text = (preview_text or "")
    name = Path(filename).stem

    con = sqlite3.connect(DB_PATH)
    cur = con.cursor()
    cur.execute("SELECT classe, marker, peso, tipo_marker FROM regole_classificazione")
    rules = cur.fetchall()
    con.close()

    scores: Dict[str, int] = {}
    trace_parts: List[str] = []

    forced: Optional[str] = None

    def add(cl: str, val: int, why: str):
        scores[cl] = scores.get(cl, 0) + int(val)
        trace_parts.append(f"{cl}:{val}:{why}")

    for cl, marker, peso, tipo_marker in rules:
        marker_l = (marker or "").lower().strip()
        if not marker_l:
            continue

        hit = False
        if marker_l in text.lower():
            hit = True
        if marker_l in name.lower():
            hit = True

        if not hit:
            continue

        if (tipo_marker or "").lower() == "blacklist":
            add(cl, int(peso), f"BLACK:{marker_l}")
        else:
            add(cl, int(peso), f"hit:{marker_l}")

    # decisione: scegli max score
    if not scores:
        return "ALTRO", 0, "no_rules_hit", None

    best_cls = max(scores.items(), key=lambda kv: kv[1])[0]
    best_score = scores[best_cls]

    # soglia minima: se score < 3 -> ALTRO
    if best_score < 3:
        return "ALTRO", best_score, " | ".join(trace_parts), None

    return best_cls, best_score, " | ".join(trace_parts), forced


def main():
    keep_last_n_files(LOG_DIR, n=5)

    LIB_DIR.mkdir(parents=True, exist_ok=True)
    (LIB_DIR / "revisione" / "occ").mkdir(parents=True, exist_ok=True)
    (LIB_DIR / "report").mkdir(parents=True, exist_ok=True)

    ensure_db()
    seed_rules_if_empty()

    files = iter_input_files()

    print(f"[INFO] DB: {DB_PATH}")
    print(f"[INFO] Input: {INPUT_DIR}")
    print(f"[INFO] File trovati: {len(files)}")
    print(f"[INFO] PDF lib: {PDF_LIB or 'NONE'}\n")

    for f in files:
        ext = f.suffix.lower()
        ocr_state = "⏭"
        rn_state = "⏭"
        mv_state = "⏭"
        note_msg = ""

        sha1 = sha1_file(f)

        # ==========================================================
        # DUPLICATI (VERI) -> input_documenti\duplicati e STOP
        # ==========================================================
        existing = fetch_doc_by_sha1(sha1)
        if existing:
            (_id, _nome, _tipo, _rgnr, _anno, _proc, _dda, _op, _nomeind, _status, _percorso, _forza,
             _tp, _snip, _mr, _rc, _man, _cs, _ct, _cf, _note) = existing
            dup_vero = is_final_status(_status) or (not is_path_in_input(_percorso))
            if dup_vero:
                try:
                    moved = move_to_duplicati(f)
                    print(f"[DUP] {f.name} -> spostato in input_documenti\\duplicati\\{moved.name} (sha1 già presente)")
                except Exception as e:
                    print(f"[DUP] {f.name} -> sha1 già presente ma move duplicati FALLITO: {e}")
                continue

        rel_path = str(Path("input_documenti") / f.name)

        # INSERT/UPDATE identity
        if existing:
            db_update_identity(existing[0], f.name, rel_path, f.stat().st_size, format_mtime(f.stat().st_mtime))
        else:
            db_insert_new_file(f, sha1, rel_path)

        row = fetch_doc_by_sha1(sha1)
        if not row:
            print(f"{f.name} | INSERT ⚠️ | OCR ⏭ | RENAME ⏭ | MOVE ⏭ | NOTE db_missing_after_insert")
            continue

        (doc_id, db_nome_file, tipo_doc, rgnr, anno, procura, dda_flag,
         operazione_nome, nome_indagine, status, percorso_file, forza_polizia,
         text_preview, snippet_testo, motivo_rev, retry_count, manual_required,
         class_score, class_trace, class_forced, note_db) = row

        # ==========================================================
        # PRE-PROCESS PER ESTENSIONE (DOC/RTF/TXT SEMPRE)
        # ==========================================================
        if ext in (".doc", ".docx", ".rtf", ".txt"):
            ok_doc, msg_doc = run_worker(Path("script/workers/organizzatore_doc_reader.py"), ["--file", db_nome_file])
            ocr_state = mark(ok_doc)
            if not ok_doc and not note_msg:
                note_msg = msg_doc

        # ==========================================================
        # CLASSIFICAZIONE RIGIDA (tabella regole) su preview + nome file
        # ==========================================================
        # ricarica preview aggiornato dopo doc_reader
        row = fetch_doc_by_sha1(sha1)
        (doc_id, db_nome_file, tipo_doc, rgnr, anno, procura, dda_flag,
         operazione_nome, nome_indagine, status, percorso_file, forza_polizia,
         text_preview, snippet_testo, motivo_rev, retry_count, manual_required,
         class_score, class_trace, class_forced, note_db) = row

        decided_tipo, score, trace, forced = score_with_rules(text_preview or "", db_nome_file)

        # aggiorna solo se cambia
        tipo_doc_new = decided_tipo
        status_new = "CLASSIFIED" if decided_tipo == "OCC" else "NEEDS_DEEPER_ANALYSIS"

        if (tipo_doc_new != (tipo_doc or "")) or (status_new != (status or "")) or (trace != (class_trace or "")) or (int(score) != int(class_score or 0)):
            db_update_classification(doc_id, tipo_doc_new, status_new, int(score), trace, forced)

        # ricarica post-update
        row = fetch_doc_by_sha1(sha1)
        (doc_id, db_nome_file, tipo_doc, rgnr, anno, procura, dda_flag,
         operazione_nome, nome_indagine, status, percorso_file, forza_polizia,
         text_preview, snippet_testo, motivo_rev, retry_count, manual_required,
         class_score, class_trace, class_forced, note_db) = row

        tipo_up = (tipo_doc or "").upper()

        # ==========================
        # NON-OCC:
        # - se nel nome c'è OCC/OCCC -> label fix rename (v16) e poi prosegue
        # - se anno noto -> percorso (libreria\YYYY)
        # - se anno assente -> vari_verbali
        # ==========================
        if tipo_up != "OCC":
            _ok_probe, _msg_probe = run_worker(Path("script/workers/organizzatore_altro_probe.py"), ["--file", db_nome_file])

            # Se NON-OCC ma nel nome c'è "OCC/OCCC", correggiamo l'etichetta col rename (v16)
            # (es: "OCC FEHIDA.doc" -> "INFORMATIVA FEHIDA.doc")
            if filename_has_occ_token(db_nome_file):
                _ok_rn_fix, _msg_rn_fix = run_worker(Path("script/workers/organizzatore_rename.py"), ["--file", db_nome_file])
                rn_state = mark(_ok_rn_fix)
                if (not _ok_rn_fix) and (not note_msg):
                    note_msg = _msg_rn_fix
                row_fix = fetch_doc_by_sha1(sha1)
                if row_fix:
                    db_nome_file = row_fix[1]
                    tipo_doc = row_fix[2]
                    anno = row_fix[4]
                    status = row_fix[9]
                    percorso_file = row_fix[10]

            if anno:
                ok_mv, msg_mv = run_worker(Path("script/workers/organizzatore_percorso.py"), ["--file", db_nome_file])
                mv_state = mark(ok_mv)
                if not ok_mv and not note_msg:
                    note_msg = msg_mv

                row_v = fetch_doc_by_sha1(sha1)
                extra = f"{fmt_tipo(row_v[2])} anno={row_v[4]} status={row_v[9]} path={row_v[10]}"
                print(f"{f.name} | OCR {ocr_state} | RENAME {rn_state} | MOVE {mv_state} | {extra}" + (f" | NOTE {note_msg}" if note_msg else ""))
                continue

            ok_vv, msg_vv = run_worker(Path("script/workers/organizzatore_vari_verbali.py"), ["--file", db_nome_file])
            mv_state = mark(ok_vv)
            if not ok_vv and not note_msg:
                note_msg = msg_vv

            row_v = fetch_doc_by_sha1(sha1)
            extra = f"{fmt_tipo(row_v[2])} score={row_v[17] or 0} forced={row_v[19]} status={row_v[9]} path={row_v[10]}"
            print(f"{f.name} | OCR {ocr_state} | RENAME {rn_state} | MOVE {mv_state} | {extra}" + (f" | NOTE {note_msg}" if note_msg else ""))
            continue

        # ==========================
        # OCC: estrazione rapida specifica (PDF)
        # ==========================
        if ext == ".pdf":
            ok, msg = run_worker(Path("script/workers/organizzatore_occ_ocr_rgnr.py"), ["--file", db_nome_file])
            if ocr_state == "⏭":
                ocr_state = mark(ok)
            if not ok and not note_msg:
                note_msg = msg

        # ricarica per gating OCC
        row2 = fetch_doc_by_sha1(sha1)
        has_rgnr = bool(row2[3] and RE_RGNR_PARSE.search(str(row2[3])))
        has_anno = bool(row2[4])
        has_proc = bool((row2[5] or "").strip())

        # Completo -> rename + move anno
        if has_rgnr and has_anno and has_proc and (row2[9] or "").upper() not in ("NEEDS_READER", "NEEDS_DEEPER_ANALYSIS", "STANDBY"):
            ok_rn, msg_rn = run_worker(Path("script/workers/organizzatore_rename.py"), ["--file", row2[1]])
            rn_state = mark(ok_rn)
            if not ok_rn and not note_msg:
                note_msg = msg_rn

            row3 = fetch_doc_by_sha1(sha1)
            ok_mv, msg_mv = run_worker(Path("script/workers/organizzatore_percorso.py"), ["--file", row3[1]])
            mv_state = mark(ok_mv)
            if not ok_mv and not note_msg:
                note_msg = msg_mv

            final = fetch_doc_by_sha1(sha1)
            extra = f"{fmt_tipo(final[2])} rgnr={final[3]} anno={final[4]} dda={final[6]} procura={final[5]} status={final[9]} path={final[10]}"
            print(f"{f.name} | OCR {ocr_state} | RENAME {rn_state} | MOVE {mv_state} | {extra}" + (f" | NOTE {note_msg}" if note_msg else ""))
            continue

        # Incompleto -> revisione OCC
        ok_rev, msg_rev = move_occ_to_revisione(f, sha1)
        mv_state = mark(ok_rev)
        if not ok_rev and not note_msg:
            note_msg = msg_rev

        rowf = fetch_doc_by_sha1(sha1)
        extra = f"{fmt_tipo(rowf[2])} rgnr={rowf[3]} anno={rowf[4]} dda={rowf[6]} procura={rowf[5]} status={rowf[9]} path={rowf[10]}"
        print(f"{f.name} | OCR {ocr_state} | RENAME {rn_state} | MOVE {mv_state} | {extra}" + (f" | NOTE {note_msg}" if note_msg else ""))

    # report
    ok_rep, msg_rep = run_worker(Path("script/workers/organizzatore_report_errori.py"), [])
    if ok_rep and msg_rep:
        print("\n[REPORT] " + msg_rep)
    elif not ok_rep:
        print("\n[REPORT] ⚠️ " + msg_rep)

    # ROUTER REVISIONE
    ok_router, msg_router = run_worker(Path("script/workers/organizzatore_revisione_router.py"), [])
    if ok_router and msg_router:
        print("\n[REVISIONE] " + msg_router)
    elif not ok_router:
        print("\n[REVISIONE] ⚠️ " + msg_router)

    print("\n[NOTE] Pipeline completata.\n")
=======
Organizzatore.py (v2 - ANNO + RGNR adiacente)
Smista documenti giudiziari per ANNO (cartelle YYYY) estraendo il "Procedimento Penale/anno"
accanto alle diciture RGNR / R.G.N.R. (o R.G.I.P., R.O.C.C., ecc.), con OCR di fallback.
"""

import os
import sys
import re
import csv
import hashlib
import logging
import sqlite3
import datetime
import shutil
from pathlib import Path

import fitz
import easyocr
import docx
from PIL import Image
import numpy as np
from chardet.universaldetector import UniversalDetector

try:
    from win32com import client as win32_client
    from pythoncom import com_error
except Exception:
    win32_client = None
    com_error = Exception


def _detect_project_root() -> Path:
    current = Path(__file__).resolve()
    for parent in [current.parent, *current.parents]:
        if parent.name.upper() == "SANGIA":
            return parent
    return current.parent


PROJECT_ROOT = _detect_project_root()
LIBRERIA_BASE_PATH = PROJECT_ROOT / "libreria"
INPUT_BASE_PATH = PROJECT_ROOT / "input_documenti"
BACKUP_BASE_PATH = PROJECT_ROOT / "Backup"
LOG_BASE_PATH = BACKUP_BASE_PATH / "log_smistamenti"

CARTELLE_VARI_VERBALI = "Vari_Verbali"
CARTELLA_ERRORI = "Errori"
CARTELLA_DUPLICATI = "Duplicati"
DB_NAME = "documenti.db"
SOTTOCARTELLA_SENTENZE = "Sentenze"

base_path = LIBRERIA_BASE_PATH
input_path = INPUT_BASE_PATH
error_path = INPUT_BASE_PATH / CARTELLA_ERRORI
varie_verbali_path = base_path / CARTELLE_VARI_VERBALI
duplicates_path = INPUT_BASE_PATH / CARTELLA_DUPLICATI
db_path = base_path / DB_NAME
error_report_path = error_path / "errori_report.csv"
manual_overrides_path = BACKUP_BASE_PATH / "correzioni" / "correzioni_smistamento.csv"

report_counts = {
    "totale": 0,
    "smistati_per_anno": 0,
    "smistati_verbali": 0,
    "errori": 0,
    "saltati_sha1": 0,
    "saltati_nome": 0,
    "duplicati_archiviati": 0,
}

CURRENT_YEAR = datetime.datetime.now().year
MANUAL_OVERRIDES = {}


def setup_logging():
    LOG_BASE_PATH.mkdir(parents=True, exist_ok=True)
    log_filename = datetime.datetime.now().strftime("log_smistamento_%Y%m%d_%H%M%S.log")
    log_path = LOG_BASE_PATH / log_filename
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s - %(levelname)s - %(message)s",
        handlers=[
            logging.FileHandler(log_path, encoding="utf-8"),
            logging.StreamHandler(sys.stdout),
        ],
    )
    logging.info("--- Avvio del Processo di Smistamento ---")


def init_ocr_reader():
    gpu_enabled = False
    gpu_reason = "torch non disponibile"
    try:
        import torch

        gpu_enabled = bool(torch.cuda.is_available())
        if gpu_enabled:
            gpu_reason = "CUDA disponibile"
        else:
            gpu_reason = "CUDA non disponibile (driver/GPU non rilevati)"
    except Exception as e:
        gpu_enabled = False
        gpu_reason = f"torch/cuda check fallito: {e}"

    try:
        reader = easyocr.Reader(["it"], gpu=gpu_enabled)
        if gpu_enabled:
            logging.info("OCR Reader inizializzato (GPU): %s", gpu_reason)
        else:
            logging.info("OCR Reader inizializzato (CPU): %s", gpu_reason)
        return reader
    except Exception as e:
        logging.warning("Init OCR con GPU=%s fallita (%s). Fallback CPU.", gpu_enabled, e)
        return easyocr.Reader(["it"], gpu=False)


def normalize_input_dropzone():
    return


def load_manual_overrides() -> dict:
    overrides = {}
    if not manual_overrides_path.exists():
        return overrides

    try:
        with open(manual_overrides_path, "r", encoding="utf-8-sig", newline="") as f:
            reader = csv.DictReader(f, delimiter=";")
            for row in reader:
                file_name = (row.get("file") or "").strip()
                if not file_name or file_name.startswith("#"):
                    continue

                anno_val = (row.get("anno") or "").strip()
                anno = None
                if anno_val:
                    try:
                        anno = int(anno_val)
                    except ValueError:
                        anno = normalize_year(anno_val)

                overrides[file_name] = {
                    "anno": anno,
                    "tipo_documento": (row.get("tipo_documento") or "").strip() or None,
                    "rgnr": (row.get("rgnr") or "").strip() or None,
                    "procura": (row.get("procura") or "").strip() or None,
                    "destinazione": (row.get("destinazione") or "").strip().upper() or None,
                    "note": (row.get("note") or "").strip() or None,
                }
        logging.info("Override manuali caricati: %s (%d righe)", manual_overrides_path, len(overrides))
    except Exception as e:
        logging.warning("Impossibile leggere override manuali %s: %s", manual_overrides_path, e)

    return overrides


def ensure_manual_overrides_template():
    if manual_overrides_path.exists():
        return

    try:
        manual_overrides_path.parent.mkdir(parents=True, exist_ok=True)
        with open(manual_overrides_path, "w", encoding="utf-8", newline="") as f:
            f.write("file;anno;tipo_documento;rgnr;procura;destinazione;note\n")
        logging.info("Creato template override manuali: %s", manual_overrides_path)
    except Exception as e:
        logging.warning("Impossibile creare template override %s: %s", manual_overrides_path, e)


def _append_error_report(file_name: str, reason: str):
    error_path.mkdir(parents=True, exist_ok=True)
    exists = error_report_path.exists()
    with open(error_report_path, "a", encoding="utf-8") as f:
        if not exists:
            f.write("timestamp;file;reason\n")
        ts = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        f.write(f"{ts};{file_name};{reason}\n")


def move_to_error(file_path: Path, reason: str):
    error_path.mkdir(parents=True, exist_ok=True)
    dest = error_path / file_path.name
    try:
        if dest.exists():
            stem = file_path.stem
            suffix = file_path.suffix
            dest = error_path / f"{stem}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}{suffix}"
        shutil.move(str(file_path), str(dest))
    except Exception as e:
        logging.error("Errore spostamento in cartella errori %s: %s", file_path.name, e)
    _append_error_report(file_path.name, reason)
    report_counts["errori"] += 1
    logging.warning("Errore classificazione %s: %s", file_path.name, reason)


def setup_database():
    conn = sqlite3.connect(db_path)
    cur = conn.cursor()
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS documenti (
            id INTEGER PRIMARY KEY,
            nome_file TEXT NOT NULL,
            rgnr TEXT,
            anno INTEGER,
            procura TEXT,
            tipo_documento TEXT,
            indagato_principale TEXT,
            num_correi INTEGER,
            operazione_nome TEXT,
            data_riferimento_file TEXT,
            modello_rgnr TEXT,
            dimensione_file INTEGER,
            data_ultima_modifica TEXT,
            sha1 TEXT,
            tipo_file TEXT,
            numero_pagine INTEGER,
            text_source TEXT,
            needs_ocr INTEGER,
            text_quality INTEGER,
            snippet_testo TEXT,
            note TEXT,
            data_inserimento TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
        """
    )
    cur.execute("CREATE INDEX IF NOT EXISTS idx_anno ON documenti(anno)")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_nome_file ON documenti(nome_file)")
    cur.execute("CREATE INDEX IF NOT EXISTS idx_sha1 ON documenti(sha1)")
    for col_def in [
        "tipo_file TEXT",
        "numero_pagine INTEGER",
        "text_source TEXT",
        "needs_ocr INTEGER",
        "text_quality INTEGER",
        "snippet_testo TEXT",
    ]:
        try:
            cur.execute(f"ALTER TABLE documenti ADD COLUMN {col_def}")
        except sqlite3.OperationalError:
            pass
    conn.commit()
    conn.close()


def insert_document_info(
    nome_file,
    rgnr,
    anno,
    procura,
    tipo_documento,
    indagato_principale,
    num_correi,
    operazione_nome,
    data_riferimento_file,
    modello_rgnr,
    dimensione,
    data_modifica,
    sha1sum,
    tipo_file,
    numero_pagine,
    text_source,
    needs_ocr,
    text_quality,
    snippet_testo,
):
    conn = sqlite3.connect(db_path)
    cur = conn.cursor()
    cur.execute(
        """
        INSERT INTO documenti (
            nome_file, rgnr, anno, procura, tipo_documento,
            indagato_principale, num_correi, operazione_nome,
            data_riferimento_file, modello_rgnr, dimensione_file,
            data_ultima_modifica, sha1, tipo_file, numero_pagine,
            text_source, needs_ocr, text_quality, snippet_testo
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (
            nome_file,
            rgnr,
            anno,
            procura,
            tipo_documento,
            indagato_principale,
            num_correi,
            operazione_nome,
            data_riferimento_file,
            modello_rgnr,
            dimensione,
            data_modifica,
            sha1sum,
            tipo_file,
            numero_pagine,
            text_source,
            needs_ocr,
            text_quality,
            snippet_testo,
        ),
    )
    conn.commit()
    conn.close()


def file_is_already_processed_by_name(file_name):
    conn = sqlite3.connect(db_path)
    cur = conn.cursor()
    cur.execute("SELECT 1 FROM documenti WHERE nome_file = ?", (file_name,))
    ok = cur.fetchone() is not None
    conn.close()
    return ok


def file_is_already_processed_by_hash(sha1sum):
    conn = sqlite3.connect(db_path)
    cur = conn.cursor()
    cur.execute("SELECT 1 FROM documenti WHERE sha1 = ?", (sha1sum,))
    ok = cur.fetchone() is not None
    conn.close()
    return ok


def file_sha1(path: Path) -> str:
    h = hashlib.sha1()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()


def extract_info_from_filename(filename):
    info = {
        "indagato_principale": None,
        "num_correi": None,
        "operazione_nome": None,
        "data_riferimento_file": None,
    }
    base_name = Path(filename).stem
    m = re.match(r"(.+?)\s*\+\s*(\d+)\s*(.*)", base_name, re.IGNORECASE)
    if m:
        indagato_raw = m.group(1).strip()
        info["num_correi"] = int(m.group(2))
        remainder = m.group(3).strip()
        info["indagato_principale"] = re.sub(r"(\s+\d+|\s+Olimpia)$", "", indagato_raw).strip()
        md = re.search(r"(\d{8}|\d{6}|\d{4}[._\-]\d{2}[._\-]\d{2})", remainder)
        if md:
            info["data_riferimento_file"] = re.sub(r"[._\-]", "", md.group(1))
        mo = re.search(r"\(?([A-Z][A-Za-z0-9\s]+)\)?", remainder)
        if mo:
            info["operazione_nome"] = mo.group(1).strip().replace("(", "").replace(")", "")
    return info


def classify_document_type(text):
    t = text.lower()
    keywords = {
        "Sentenza d'Appello": ["sentenza d'appello"],
        "Ordinanza di Custodia Cautelare in Carcere": [
            "ordinanza di custodia cautelare in carcere",
            "occ",
            "ordinanza di custodia cautelare",
        ],
        "Ordinanza e Sequestro Preventivo": [
            "decreto di sequestro preventivo",
            "sequestro preventivo",
        ],
        "Ordinanza": ["ordinanza"],
        "Richiesta di Misura Cautelare": [
            "richiesta di custodia cautelare",
            "richiesta di misura cautelare",
        ],
        "Annotazione di Polizia Giudiziaria": [
            "annotazione di polizia giudiziaria",
            "annotazione p.g.",
            "annotazione",
        ],
        "Relazione di Servizio": ["relazione di servizio", "relazione"],
        "Articolo di Giornale": ["articolo di giornale", "rassegna stampa", "articolo"],
        "Informativa": ["informativa"],
        "Dichiarazioni": ["dichiarazioni", "verbale"],
        "Fascicolo Personale": ["fascicolo personale", "fascicolo"],
        "Sentenza": ["sentenza", "sentenze"],
    }
    for doc_type, k_list in keywords.items():
        for k in k_list:
            if re.search(r"\b" + re.escape(k) + r"\b", t):
                return doc_type
    return "Non Classificato"


def detect_file_type(file_path: Path) -> str:
    ext = file_path.suffix.lower()
    if ext == ".pdf":
        return "pdf"
    if ext in {".docx", ".doc"}:
        return "docx" if ext == ".docx" else "doc"
    if ext in {".txt", ".csv", ".json", ".xml"}:
        return "plain"
    if ext == ".rtf":
        return "rtf"
    if ext in {".jpg", ".jpeg", ".png", ".tif", ".tiff", ".bmp", ".webp"}:
        return "image"

    try:
        with open(file_path, "rb") as f:
            head = f.read(16)
        if head.startswith(b"%PDF"):
            return "pdf"
        if head.startswith(b"{\\rtf"):
            return "rtf"
        if head[:4] == b"PK\x03\x04":
            return "docx"
    except Exception:
        pass

    return "unknown"


def _text_quality_score(text: str) -> int:
    t = (text or "").strip()
    if not t:
        return 0
    alpha = sum(c.isalpha() for c in t)
    printable = sum(c.isprintable() for c in t)
    ratio_alpha = alpha / max(1, len(t))
    ratio_print = printable / max(1, len(t))
    len_score = min(60, len(t) // 30)
    return max(0, min(100, int(len_score + ratio_alpha * 30 + ratio_print * 10)))


def _read_plain_with_detection(file_path: Path, max_chars: int) -> str:
    try:
        return file_path.read_text(encoding="utf-8")[:max_chars]
    except Exception:
        det = UniversalDetector()
        with open(file_path, "rb") as f:
            for line in f:
                det.feed(line)
                if det.done:
                    break
        det.close()
        enc = det.result.get("encoding") or "latin-1"
        return file_path.read_text(encoding=enc, errors="ignore")[:max_chars]


def _rtf_to_text(rtf_text: str) -> str:
    t = re.sub(r"\\par[d]?", "\n", rtf_text)
    t = re.sub(r"\\'[0-9a-fA-F]{2}", "", t)
    t = re.sub(r"\\[a-zA-Z]+-?\d* ?", "", t)
    t = t.replace("{", " ").replace("}", " ")
    return re.sub(r"\s+", " ", t).strip()


def _ocr_image_np(img_np, ocr_reader):
    if ocr_reader is None:
        return ""
    out = []
    try:
        for (_, txt, prob) in ocr_reader.readtext(img_np):
            if prob >= 0.3:
                out.append(txt)
    except Exception:
        return ""
    return " ".join(out)


def get_text_from_file(file_path: Path, ocr_reader):
    MAX_CHARS_FOR_SEARCH = 5000
    text_content = ""
    moved_to_error = False
    file_type = detect_file_type(file_path)
    meta = {
        "tipo_file": file_type,
        "text_source": "none",
        "needs_ocr": 0,
        "text_quality": 0,
        "numero_pagine": None,
        "snippet_testo": "",
        "ocr_sample_only": 0,
    }

    if file_type == "docx":
        try:
            d = docx.Document(file_path)
            parts = [p.text for p in d.paragraphs if p.text]
            for t in d.tables:
                for row in t.rows:
                    for cell in row.cells:
                        if cell.text:
                            parts.append(cell.text)
            text_content = "\n".join(parts)[:MAX_CHARS_FOR_SEARCH]
            meta["text_source"] = "docx"
        except Exception as e:
            logging.error(f"Errore lettura DOCX: {file_path.name}. Dettaglio: {e}")

    elif file_type == "doc":
        if win32_client is None:
            logging.error(f"pywin32 non disponibile: impossibile leggere DOC {file_path.name}")
        else:
            word = None
            try:
                word = win32_client.Dispatch("Word.Application")
                doc = word.Documents.Open(str(file_path))
                text_content = (doc.Range(0, MAX_CHARS_FOR_SEARCH).Text or "")
                doc.Close(SaveChanges=0)
                meta["text_source"] = "docx"
            except com_error as e:
                logging.error(f"Errore COM DOC: {file_path.name}. Dettaglio: {e}")
                try:
                    move_to_error(file_path, "DOC corrotto o non leggibile via COM")
                    moved_to_error = True
                except Exception as move_e:
                    logging.critical(f"Impossibile spostare DOC corrotto: {move_e}")
            except Exception as e:
                logging.error(f"Errore generico DOC: {file_path.name}. Dettaglio: {e}")
            finally:
                if word:
                    try:
                        word.Quit()
                    except Exception:
                        pass

    elif file_type == "plain":
        try:
            text_content = _read_plain_with_detection(file_path, MAX_CHARS_FOR_SEARCH)
            meta["text_source"] = "plain"
        except Exception as e:
            logging.error(f"Errore lettura testo: {file_path.name}. Dettaglio: {e}")

    elif file_type == "rtf":
        try:
            raw = _read_plain_with_detection(file_path, max_chars=MAX_CHARS_FOR_SEARCH * 2)
            text_content = _rtf_to_text(raw)[:MAX_CHARS_FOR_SEARCH]
            meta["text_source"] = "plain"
        except Exception as e:
            logging.error(f"Errore lettura RTF: {file_path.name}. Dettaglio: {e}")

    elif file_type == "image":
        meta["needs_ocr"] = 1
        try:
            img = Image.open(file_path).convert("RGB")
            text_content = _ocr_image_np(np.array(img), ocr_reader)[:MAX_CHARS_FOR_SEARCH]
            meta["text_source"] = "ocr"
            meta["numero_pagine"] = 1
        except Exception as e:
            logging.error(f"Errore OCR immagine: {file_path.name}. Dettaglio: {e}")

    elif file_type == "pdf":
        try:
            with fitz.open(file_path) as doc:
                meta["numero_pagine"] = doc.page_count
                digital_parts = []
                for i in range(doc.page_count):
                    digital_parts.append(doc.load_page(i).get_text() or "")
                    if len("".join(digital_parts)) >= MAX_CHARS_FOR_SEARCH:
                        break
                digital_text = " ".join(digital_parts)[:MAX_CHARS_FOR_SEARCH]
                if _text_quality_score(digital_text) >= 20 and len(digital_text.strip()) >= 120:
                    text_content = digital_text
                    meta["text_source"] = "pdf_text"
                else:
                    meta["needs_ocr"] = 1
                    indexes = set(range(min(3, doc.page_count)))
                    if doc.page_count > 0:
                        indexes.add(doc.page_count - 1)
                    if doc.page_count > 2:
                        indexes.add(doc.page_count // 2)
                    sample_parts = []
                    for i in sorted(indexes):
                        page = doc.load_page(i)
                        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
                        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                        sample_parts.append(_ocr_image_np(np.array(img), ocr_reader))
                        if len(" ".join(sample_parts)) >= MAX_CHARS_FOR_SEARCH:
                            break
                    text_content = " ".join(sample_parts)[:MAX_CHARS_FOR_SEARCH]
                    meta["text_source"] = "ocr"
                    meta["ocr_sample_only"] = 1
        except Exception as e:
            logging.error(f"Errore lettura PDF: {file_path.name}. Dettaglio: {e}")

    else:
        try:
            text_content = _read_plain_with_detection(file_path, MAX_CHARS_FOR_SEARCH)
            meta["text_source"] = "plain"
        except Exception:
            text_content = ""
            meta["text_source"] = "none"

    text_content = (text_content or "")[:MAX_CHARS_FOR_SEARCH]
    meta["text_quality"] = _text_quality_score(text_content)
    meta["snippet_testo"] = text_content[:280]
    return text_content, moved_to_error, meta


def normalize_year(two_or_four: str):
    try:
        n = int(two_or_four)
    except ValueError:
        return None
    if 0 <= n <= 99:
        if 80 <= n <= 99:
            return 1900 + n
        pivot = CURRENT_YEAR % 100
        return 2000 + n if n <= pivot else 1900 + n
    if 1900 <= n <= CURRENT_YEAR + 1:
        return n
    return None


def find_near_pairs(text_upper: str, token_spans, window: int = 60):
    pairs = []
    pp_pat = re.compile(r"(?:N\.?\s*)?(\d{1,6})\s*[/\-]\s*(\d{2}|\d{4})")
    for s, e in token_spans:
        left = max(0, s - window)
        right = min(len(text_upper), e + window)
        for m in pp_pat.finditer(text_upper[left:right]):
            pairs.append((left + m.start(), m.group(1), m.group(2)))

    ranked = []
    for abs_start, num, yr in pairs:
        dists = [abs(abs_start - ((s + e) // 2)) for s, e in token_spans]
        ranked.append((min(dists), abs_start, num, yr))
    ranked.sort(key=lambda x: x[0])
    return [(a, b, c) for _, a, b, c in ranked]


def find_year_and_info_in_text(text: str):
    info = {
        "year": None,
        "rgnr": None,
        "procura": None,
        "modello_rgnr": None,
        "is_comune_ordinance": False,
    }

    text_upper = text.upper().replace("\n", " ")
    mp = re.search(r"(?:TRIBUNALE|PROCURA|CORTE|LEGIONE CARABINIERI|D\.?D\.?A\.?|COMUNE)\s+DI\s+([A-Z\s\(\)]+)", text_upper)
    if mp:
        info["procura"] = mp.group(1).strip()

    if re.search(r"\bCOMUNE\s+DI\b", text_upper) and not re.search(
        r"\b(TRIBUNALE|PROCURA|R\.?\s*G\.?\s*N\.?\s*R\.?|RGNR|R\.?\s*G\.?\s*I\.?\s*P\.?)\b",
        text_upper,
    ):
        info["is_comune_ordinance"] = True

    candidates = []
    weighted_patterns = [
        (1, r"\bN\.?\s*(\d{1,6})\s*[/\-]\s*(\d{2,4})\s*R\.?\s*G\.?\s*N\.?\s*R\.?\s*D\.?\s*D\.?\s*A\.?"),
        (2, r"\bN\.?\s*(\d{1,6})\s*[/\-]\s*(\d{2,4})\s*(?:R\.?\s*G\.?\s*N\.?\s*R\.?|RGNR)"),
        (3, r"\bN\.?\s*(\d{1,6})\s*[/\-]\s*(\d{2,4})\s*R\.?\s*G\.?\s*NOTIZIE\s+DI\s+REATO"),
        (4, r"\b(\d{1,6})\s*[/\-]\s*(\d{2,4})[^\n]{0,35}(?:R\.?\s*G\.?\s*N\.?\s*R\.?|RGNR)"),
        (6, r"PROC\.?\s*N\.?\s*(\d{1,6})\s*[/\-]\s*(\d{2,4})\s*(?:R\.?\s*G\.?\s*N\.?\s*R\.?|R\.?\s*G\.?\s*I\.?\s*P\.?|R\.?\s*O\.?\s*C\.?\s*C\.?)"),
        (7, r"\bN\.?\s*(\d{1,6})\s*[/\-]\s*(\d{2,4})\s*(?:R\.?\s*O\.?\s*C\.?\s*C\.?)"),
        (8, r"\b(\d{1,6})\s*[/\-]\s*(\d{2,4})\s*R\.?\s*G\.?\b"),
    ]

    for weight, pat in weighted_patterns:
        for m in re.finditer(pat, text_upper):
            yr = normalize_year(m.group(2))
            if yr:
                candidates.append((weight, m.start(), m.group(1), m.group(2), yr))

    token_regex = re.compile(
        r"(R\.?\s*G\.?\s*N\.?\s*R\.?|RGNR|R\.?\s*G\.?\s*I\.?\s*P\.?|R\.?\s*O\.?\s*C\.?\s*C\.?|R\.?\s*G\.?)"
    )
    token_spans = [m.span() for m in token_regex.finditer(text_upper)]
    for _, num_str, year_str in find_near_pairs(text_upper, token_spans, window=60):
        yr = normalize_year(year_str)
        if yr:
            near_pos = text_upper.find(f"{num_str}/{year_str}")
            candidates.append((9, near_pos if near_pos >= 0 else 999999, num_str, year_str, yr))

    if candidates:
        candidates.sort(key=lambda x: (x[0], x[1]))
        _, _, num_raw, year_raw, year_norm = candidates[0]
        info["year"] = year_norm
        info["rgnr"] = f"{num_raw}/{year_raw}"

    if info["year"] is None:
        m = re.search(r"(\d{1,6})\s*[/\-]\s*(\d{2,4})", text_upper)
        if m:
            yr = normalize_year(m.group(2))
            if yr:
                info["year"] = yr
                info["rgnr"] = f"{m.group(1)}/{m.group(2)}"

    if info["rgnr"]:
        idx = text_upper.find(info["rgnr"])
        if idx != -1:
            mm = re.search(r"MOD\.?(?:ELLO)?[\s._\-]*(\d{2})", text_upper[idx : idx + 120])
            if mm and mm.group(1) in {"21", "44"}:
                info["modello_rgnr"] = f"Mod. {mm.group(1)}"

    return info


def extract_year_from_text_dates(text: str):
    t = text.upper()

    for m in re.finditer(r"\b(\d{1,2})[./\-](\d{1,2})[./\-](\d{2,4})\b", t):
        yr = normalize_year(m.group(3))
        if yr and 1980 <= yr <= CURRENT_YEAR + 1:
            return yr

    months = "GENNAIO|FEBBRAIO|MARZO|APRILE|MAGGIO|GIUGNO|LUGLIO|AGOSTO|SETTEMBRE|OTTOBRE|NOVEMBRE|DICEMBRE"
    m = re.search(rf"\b\d{{1,2}}\s+(?:{months})\s+(\d{{4}})\b", t)
    if m:
        yr = normalize_year(m.group(1))
        if yr and 1980 <= yr <= CURRENT_YEAR + 1:
            return yr

    return None


def build_dest_dir_by_year(base: Path, anno: int) -> Path:
    return base / str(anno)


def archive_duplicate_file(file_path: Path, reason: str):
    duplicates_path.mkdir(parents=True, exist_ok=True)
    destination = duplicates_path / file_path.name
    if destination.exists():
        stem = destination.stem
        suffix = destination.suffix
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        destination = duplicates_path / f"{stem}__dup_{timestamp}{suffix}"

    try:
        shutil.move(str(file_path), str(destination))
        report_counts["duplicati_archiviati"] += 1
        logging.info("Duplicato archiviato in %s (%s): %s", duplicates_path, reason, file_path.name)
    except Exception as e:
        logging.warning("Impossibile archiviare duplicato %s (%s): %s", file_path.name, reason, e)


def infer_year_from_filename(filename: str):
    stem = Path(filename).stem
    for m in re.finditer(r"(?<!\d)(\d{6}|\d{8})(?!\d)", stem):
        token = m.group(1)
        year_token = token[0:4] if len(token) == 8 else token[0:2]
        yr = normalize_year(year_token)
        if yr and 1980 <= yr <= CURRENT_YEAR + 1:
            return yr

    sep_patterns = [
        r"(?<!\d)(\d{2})[._/-](\d{2})[._/-](\d{2}|\d{4})(?!\d)",
        r"(?<!\d)(\d{4})[._/-](\d{2})[._/-](\d{2})(?!\d)",
    ]
    for pat in sep_patterns:
        for m in re.finditer(pat, stem):
            ypart = m.group(3) if pat.startswith(r"(?<!\d)(\d{2})") else m.group(1)
            yr = normalize_year(ypart)
            if yr and 1980 <= yr <= CURRENT_YEAR + 1:
                return yr

    return None


def parse_title_hints(filename: str) -> dict:
    stem = Path(filename).stem
    t = stem.upper()

    hints = {
        "document_type": None,
        "year": None,
        "rgnr": None,
    }

    if re.search(r"\bSENTENZA\b", t):
        hints["document_type"] = "Sentenza"

    if "ANNOTAZIONE" in t:
        hints["document_type"] = "Annotazione di Polizia Giudiziaria"
    elif "RELAZIONE" in t:
        hints["document_type"] = "Relazione di Servizio"
    elif "ARTICOLO" in t:
        hints["document_type"] = "Articolo di Giornale"

    if re.search(r"\b(OCC|ORDINANZA)\b", t):
        hints["document_type"] = "Ordinanza di Custodia Cautelare in Carcere"

    if re.search(r"\bPARTE\s*[_\- ]?(I|II|III|IV|V|1|2|3|4|5)\b", t):
        hints["document_type"] = hints["document_type"] or "Ordinanza di Custodia Cautelare in Carcere"

    if re.search(r"[A-ZÀ-ÖØ-Ý][A-ZÀ-ÖØ-Ýa-zà-öø-ÿ'`]+\s+[A-ZÀ-ÖØ-Ý][A-ZÀ-ÖØ-Ýa-zà-öø-ÿ'`]+\s*\+\s*\d+", stem):
        hints["document_type"] = hints["document_type"] or "Ordinanza di Custodia Cautelare in Carcere"

    strong_patterns = [
        r"(?:SENTENZA|SENT\.?|OCC|R\.?G\.?N\.?R\.?)[^\d]{0,20}(\d{1,6})\s*[/\-]\s*(\d{2,4})",
        r"\bN\.?\s*(\d{1,6})\s*[/\-]\s*(\d{2,4})",
    ]
    for pat in strong_patterns:
        m = re.search(pat, t, re.IGNORECASE)
        if m:
            yy = normalize_year(m.group(2))
            if yy:
                hints["year"] = yy
                hints["rgnr"] = f"{m.group(1)}/{m.group(2)}"
                break

    if hints["year"] is None:
        hints["year"] = infer_year_from_filename(filename)

    return hints


def build_no_rule_reason(file_name: str, document_type: str, extracted_info: dict, filename_info: dict, title_hints: dict, final_year):
    year_from_text = extracted_info.get("year")
    rgnr = extracted_info.get("rgnr")
    procura = extracted_info.get("procura")
    is_comune = extracted_info.get("is_comune_ordinance", False)
    data_rif = filename_info.get("data_riferimento_file")
    operazione = filename_info.get("operazione_nome")

    title_year = title_hints.get("year") if title_hints else None
    title_rgnr = title_hints.get("rgnr") if title_hints else None

    parts = [
        "Nessuna regola di smistamento applicabile",
        f"tipo_documento={document_type}",
        f"anno_finale={final_year}",
        f"anno_da_testo={year_from_text}",
        f"rgnr={rgnr}",
        f"procura={procura}",
        f"comune_ordinanza={is_comune}",
        f"data_filename={data_rif}",
        f"operazione_filename={operazione}",
        f"anno_da_titolo={title_year}",
        f"rgnr_da_titolo={title_rgnr}",
    ]

    if document_type == "Non Classificato":
        parts.append("hint=nessuna keyword documento riconosciuta")
    if final_year is None:
        parts.append("hint=anno_non_rilevato")

    return " | ".join(str(x) for x in parts)


def process_file(file_path: Path, ocr_reader):
    global report_counts
    file_name = file_path.name

    try:
        sha1sum = file_sha1(file_path)
        if file_is_already_processed_by_hash(sha1sum):
            report_counts["saltati_sha1"] += 1
            logging.info(f"File già presente (SHA1): {file_name}. Saltato.")
            archive_duplicate_file(file_path, "SHA1")
            return
    except Exception as e:
        logging.warning(f"Impossibile calcolare SHA1 per {file_name}: {e}")
        sha1sum = None

    if file_is_already_processed_by_name(file_name):
        report_counts["saltati_nome"] += 1
        logging.info(f"File già registrato per nome: {file_name}. Saltato.")
        archive_duplicate_file(file_path, "NOME")
        return

    filename_info = extract_info_from_filename(file_name)
    title_hints = parse_title_hints(file_name)
    manual_override = MANUAL_OVERRIDES.get(file_name)
    text_content, moved_to_error, extraction_meta = get_text_from_file(file_path, ocr_reader)
    if moved_to_error:
        return

    dimensione = os.path.getsize(file_path)
    data_modifica = datetime.datetime.fromtimestamp(os.path.getmtime(file_path)).strftime("%Y-%m-%d %H:%M:%S")

    if extraction_meta.get("ocr_sample_only") == 1 and extraction_meta.get("text_quality", 0) < 15:
        move_to_error(file_path, "PDF immagine: OCR campionato insufficiente (serve OCR completo manuale)")
        return

    if not text_content or len(text_content.strip()) < 10:
        move_to_error(file_path, f"Testo non leggibile o OCR insufficiente | chars={len((text_content or '').strip())}")
        return

    extracted_info = find_year_and_info_in_text(text_content)
    found_year = extracted_info["year"]
    document_type = classify_document_type(text_content)
    if title_hints.get("document_type") is not None:
        document_type = title_hints["document_type"]

    if manual_override and manual_override.get("tipo_documento"):
        document_type = manual_override["tipo_documento"]

    temp_year = None
    data_rif = filename_info.get("data_riferimento_file")
    if found_year is None and data_rif:
        year_str = data_rif[-4:] if len(data_rif) >= 8 else data_rif[-2:]
        yr = normalize_year(year_str)
        if yr and 1980 <= yr <= CURRENT_YEAR + 1:
            temp_year = yr

    final_year = found_year if found_year is not None else temp_year

    if final_year is None and title_hints.get("year") is not None:
        final_year = title_hints["year"]

    if final_year is None:
        m = re.search(r"\b(19[8-9]\d|20\d{2}|2100)\b", file_name)
        if m:
            yy = int(m.group(1))
            if 1980 <= yy <= CURRENT_YEAR + 1:
                final_year = yy

    if final_year is None:
        final_year = infer_year_from_filename(file_name)

    if final_year is None and extracted_info.get("rgnr"):
        mm = re.search(r"/(\d{2,4})", extracted_info.get("rgnr", ""))
        if mm:
            final_year = normalize_year(mm.group(1))

    if final_year is None:
        final_year = extract_year_from_text_dates(text_content)

    if manual_override and manual_override.get("anno"):
        final_year = manual_override["anno"]

    if manual_override and manual_override.get("rgnr"):
        extracted_info["rgnr"] = manual_override["rgnr"]
    if manual_override and manual_override.get("procura"):
        extracted_info["procura"] = manual_override["procura"]

    destination_override = (manual_override or {}).get("destinazione")

    target_path = None
    if destination_override == "ERRORI":
        move_to_error(file_path, f"Forzato da correzioni_smistamento.csv: {(manual_override or {}).get('note') or 'nessuna nota'}")
        return
    if final_year:
        dest_root = build_dest_dir_by_year(base_path, final_year)
        dest_root.mkdir(parents=True, exist_ok=True)
        target_path = dest_root
        report_counts["smistati_per_anno"] += 1
    else:
        move_to_error(file_path, build_no_rule_reason(file_name, document_type, extracted_info, filename_info, title_hints, final_year))
        return

    try:
        if manual_override:
            logging.info("Override manuale applicato a %s: %s", file_name, manual_override)
        shutil.move(str(file_path), str(target_path / file_name))
        if not extracted_info.get("rgnr") and title_hints.get("rgnr"):
            extracted_info["rgnr"] = title_hints["rgnr"]

        insert_document_info(
            nome_file=file_name,
            rgnr=extracted_info.get("rgnr"),
            anno=final_year,
            procura=extracted_info.get("procura"),
            tipo_documento=document_type,
            indagato_principale=filename_info.get("indagato_principale"),
            num_correi=filename_info.get("num_correi"),
            operazione_nome=filename_info.get("operazione_nome"),
            data_riferimento_file=filename_info.get("data_riferimento_file"),
            modello_rgnr=extracted_info.get("modello_rgnr"),
            dimensione=dimensione,
            data_modifica=data_modifica,
            sha1sum=sha1sum,
            tipo_file=extraction_meta.get("tipo_file"),
            numero_pagine=extraction_meta.get("numero_pagine"),
            text_source=extraction_meta.get("text_source"),
            needs_ocr=extraction_meta.get("needs_ocr"),
            text_quality=extraction_meta.get("text_quality"),
            snippet_testo=extraction_meta.get("snippet_testo"),
        )
        logging.info(f"SMISTATO: {file_name} -> {target_path}")
    except Exception as e:
        logging.error(f"ERRORE FINALE (DB/SPOSTAMENTO) {file_name}: {e}")
        if file_path.exists():
            move_to_error(file_path, f"Errore finale DB/Spostamento: {e}")
        report_counts["smistati_per_anno"] -= 1


def main():
    global MANUAL_OVERRIDES

    setup_logging()
    ocr_reader = init_ocr_reader()

    base_path.mkdir(parents=True, exist_ok=True)
    input_path.mkdir(parents=True, exist_ok=True)
    error_path.mkdir(parents=True, exist_ok=True)
    varie_verbali_path.mkdir(parents=True, exist_ok=True)
    setup_database()
    ensure_manual_overrides_template()
    MANUAL_OVERRIDES = load_manual_overrides()

    files_to_process = [f for f in input_path.iterdir() if f.is_file()]
    report_counts["totale"] = len(files_to_process)

    if report_counts["totale"] == 0:
        logging.info(f"Nessun file trovato da processare in {input_path}")
        print(f"\n--- Report Finale ---\nNessun file trovato nella cartella {input_path}")
        print("Suggerimento: inserisci i file direttamente in input_documenti")
        return

    for file_path in files_to_process:
        if file_path.exists():
            process_file(file_path, ocr_reader)

    logging.info("--- Processo Completato ---")
    print("\n--- Report Finale ---")
    print(f"File Totali Considerati: {report_counts['totale']}")
    print(f"Smistati in cartelle anno: {report_counts['smistati_per_anno']}")
    print(f"File in cartella Errori: {report_counts['errori']}")
    print(f"Report errori dettagliato: {error_report_path}")
    print("Ogni riga include il motivo tecnico della mancata classificazione.")
    print(f"File saltati per duplicato SHA1: {report_counts['saltati_sha1']}")
    print(f"File saltati per duplicato nome: {report_counts['saltati_nome']}")
    print(f"File duplicati archiviati: {report_counts['duplicati_archiviati']}")
main


if __name__ == "__main__":
    main()
